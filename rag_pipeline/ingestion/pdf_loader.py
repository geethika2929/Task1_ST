from pathlib import Path
from typing import List, Iterable
import re
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from rag_pipeline.models import Document


class DocumentIngestor:
    """
    Loads text content from .pdf, .docx, and .txt files.

    - load_all_documents(folder): scan entire folder (legacy / first build)
    - load_documents_from_paths(paths): load only given file paths (for incremental upload)
    """

    def load_all_documents(self, folder: Path) -> List[Document]:
        docs: List[Document] = []
        for path in folder.iterdir():
            if not path.is_file():
                continue
            doc = self._load_single(path)
            if doc:
                docs.append(doc)
        return docs

    def load_documents_from_paths(self, paths: Iterable[Path]) -> List[Document]:
        docs: List[Document] = []
        for path in paths:
            path = Path(path)
            if not path.is_file():
                continue
            doc = self._load_single(path)
            if doc:
                docs.append(doc)
        return docs

    def _load_single(self, path: Path) -> Document | None:
        ext = path.suffix.lower()

        if ext == ".pdf":
            text = self._extract_text_from_pdf(path)
        elif ext == ".docx":
            text = self._extract_text_from_docx(path)
        elif ext == ".txt":
            text = self._extract_text_from_txt(path)
        else:
            return None  # unsupported

        clean = self._clean_whitespace(text)

        return Document(
            doc_id=path.name,
            source_path=str(path.resolve()),
            text=clean,
        )

    def _extract_text_from_pdf(self, pdf_path: Path) -> str:
        reader = PdfReader(str(pdf_path))
        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        return "\n".join(pages_text)

    def _extract_text_from_docx(self, docx_path: Path) -> str:
        d = DocxDocument(str(docx_path))
        paras = []
        for p in d.paragraphs:
            paras.append(p.text or "")
        return "\n".join(paras)

    def _extract_text_from_txt(self, txt_path: Path) -> str:
        return txt_path.read_text(encoding="utf-8", errors="ignore")

    def _clean_whitespace(self, text: str) -> str:
        text = re.sub(r"\r", " ", text)
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()
